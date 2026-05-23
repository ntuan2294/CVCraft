"""
Template Renderer Agent - Layer 4.

Hỗ trợ 2 kiểu template:
A) Placeholder template: {{FIELD_NAME}} - dùng cho user kỹ thuật
B) Exemplar template: section header + nội dung mẫu

Auto-detect: nếu thấy {{...}} → kiểu A, không thì kiểu B.
"""
import base64
import os
import re
import copy
import tempfile
import unicodedata
from pathlib import Path
from typing import Optional
from pydantic import BaseModel, Field
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from cvcraft.generate_cv.core.state import CVAgentState
from cvcraft.generate_cv.infrastructure.llm.factory import LLMFactory, call_with_structured_output
from cvcraft.config.settings import settings


PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Z_][A-Z0-9_]*)\}\}')
KNOWN_TEMPLATE_IDS = {"1", "2", "3", "4", "5"}
PLACEHOLDER_FIELD_MAP = {
    "PHOTO": "photo",
    "IMAGE": "photo",
    "AVATAR": "photo",
    "FULL_NAME": "full_name",
    "FULLNAME": "full_name",
    "NAME": "full_name",
    "HO_TEN": "full_name",
    "JOB_TITLE": "job_title",
    "POSITION": "job_title",
    "EMAIL": "email",
    "PHONE": "phone",
    "LOCATION": "location",
    "ADDRESS": "location",
    "LINKEDIN": "linkedin",
    "GITHUB": "github",
    "CONTACT": "contact_line",
    "CONTACT_LINE": "contact_line",
    "SUMMARY": "summary",
    "PROFILE": "summary",
    "ABOUT": "summary",
    "ABOUT_ME": "summary",
    "PERSONAL_SUMMARY": "summary",
    "EXPERIENCE": "experience_section",
    "WORK_EXPERIENCE": "experience_section",
    "EXPERIENCE_SECTION": "experience_section",
    "COMPANY": "UNKNOWN",
    "TIME": "UNKNOWN",
    "EDUCATION": "education_section",
    "EDUCATION_SECTION": "education_section",
    "SCHOOL": "UNKNOWN",
    "SKILLS": "skills_section",
    "SKILLS_SECTION": "skills_section",
    "LANGUAGE": "language_section",
    "LANGUAGES": "language_section",
    "LANGUAGE_SECTION": "language_section",
    "REFERENCE": "reference_section",
    "REFERENCES": "reference_section",
    "REFERENCE_SECTION": "reference_section",
}


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            yield from nested_cell.paragraphs


def build_full_name(state: CVAgentState) -> str:
    if not state.user_profile:
        return ""
    name = state.user_profile.full_name
    if state.user_input.get("output_language") == "en":
        name = strip_accents(name)
    return compact_text(name)


def strip_accents(value: str) -> str:
    normalized = unicodedata.normalize("NFD", value)
    without_marks = "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")
    return without_marks.replace("Đ", "D").replace("đ", "d")


def compact_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def compact_lines(values: list[str]) -> list[str]:
    return [text for text in (compact_text(value) for value in values) if text]


def compact_tabbed_text(value: str) -> str:
    return "\t".join(compact_text(part) for part in (value or "").split("\t"))


def build_job_title(state: CVAgentState) -> str:
    if state.user_profile and state.user_profile.job_title:
        return state.user_profile.job_title
    if state.job_requirement and state.job_requirement.job_title:
        return state.job_requirement.job_title
    if state.user_profile and state.user_profile.work_experiences:
        return state.user_profile.work_experiences[0].position
    return ""


def build_contact_line(state: CVAgentState) -> str:
    if not state.user_profile:
        return ""
    p = state.user_profile
    parts = [compact_text(p.location or ""), compact_text(p.email), compact_text(p.phone or "")]
    parts = [x for x in parts if x]
    return " • ".join(parts)


def build_simple_field(state: CVAgentState, field: str) -> str:
    if not state.user_profile:
        return ""
    return getattr(state.user_profile, field, "") or ""


def build_summary(state: CVAgentState) -> str:
    return compact_text(state.cv_draft.summary) if state.cv_draft and state.cv_draft.summary else ""


def build_languages(state: CVAgentState) -> str:
    if not state.user_profile or not state.user_profile.languages:
        return ""
    return ", ".join(state.user_profile.languages)


def build_references(state: CVAgentState) -> str:
    if not state.user_profile:
        return ""
    fallback = "Available upon request" if state.user_input.get("output_language") == "en" else "Cung cấp khi được yêu cầu"
    return state.user_profile.references or fallback


def build_photo_filename(state: CVAgentState) -> str:
    if not state.user_profile or not state.user_profile.photo:
        return ""
    return state.user_profile.photo.get("file_name", "")


def insert_photo_at_placeholder(doc: Document, state: CVAgentState, placeholder: str) -> bool:
    if not state.user_profile or not state.user_profile.photo:
        return False

    data_url = state.user_profile.photo.get("data_url", "")
    if "," not in data_url:
        return False

    suffix = Path(state.user_profile.photo.get("file_name", "photo.png")).suffix or ".png"
    image_bytes = base64.b64decode(data_url.split(",", 1)[1])

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(image_bytes)
        tmp_path = tmp.name

    try:
        for para in iter_all_paragraphs(doc):
            if placeholder not in para.text:
                continue
            clear_paragraph_text(para)
            run = para.runs[0] if para.runs else para.add_run("")
            run.add_picture(tmp_path, width=Inches(1.15))
            return True
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    return False


def build_experience_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.experiences:
        return []

    lines = []
    present = "Present" if state.user_input.get("output_language") == "en" else "Hiện tại"
    for exp in state.cv_draft.experiences:
        end = exp.end_date or present
        lines.append(("header_with_date", f"{compact_text(exp.company).upper()}\t{compact_text(exp.start_date)} - {compact_text(end)}"))
        lines.append(("subheader", compact_text(exp.position)))
        for bullet in compact_lines(exp.bullets):
            lines.append(("bullet", bullet))
    return lines


def build_skills_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.skills_categorized:
        return []
    return [("header", f"{cat}: {', '.join(skills)}")
            for cat, skills in state.cv_draft.skills_categorized.items()]


def build_education_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.educations:
        return []

    lines = []
    present = "Present" if state.user_input.get("output_language") == "en" else "Hiện tại"
    major_joiner = " in " if state.user_input.get("output_language") == "en" else " - "
    for edu in state.cv_draft.educations:
        end = edu.end_date or present
        lines.append(("header_with_date", f"{compact_text(edu.school).upper()}\t{compact_text(edu.start_date)} - {compact_text(end)}"))
        degree_parts = [compact_text(edu.degree), compact_text(edu.major)]
        lines.append(("subheader", major_joiner.join(part for part in degree_parts if part)))
        if edu.gpa:
            lines.append(("bullet", f"GPA: {edu.gpa}"))
        for ach in compact_lines(edu.achievements):
            lines.append(("bullet", ach))
    return lines


def build_projects_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.projects:
        return []
    lines = []
    for proj in state.cv_draft.projects:
        if isinstance(proj, dict):
            name = proj.get("name", "Unnamed")
            desc = proj.get("description", "")
            lines.append(("header", name))
            if desc:
                lines.append(("bullet", desc))
            lines.append(("blank", ""))
    return lines


def build_leadership_lines(state: CVAgentState) -> list[tuple]:
    if state.cv_draft and state.cv_draft.projects:
        return build_projects_lines(state)
    return []


def build_language_lines(state: CVAgentState) -> list[tuple]:
    if not state.user_profile or not state.user_profile.languages:
        return []
    return [("bullet", item) for item in state.user_profile.languages]


def build_reference_lines(state: CVAgentState) -> list[tuple]:
    value = build_references(state)
    return [("bullet", value)] if value else []


def _date_range(start: str, end: Optional[str], state: CVAgentState) -> str:
    present = "Present" if state.user_input.get("output_language") == "en" else "Hiện tại"
    return f"{compact_text(start)} - {compact_text(end or present)}".strip(" -")


def build_time_placeholder(start: str, end: Optional[str], state: CVAgentState) -> str:
    return _date_range(start, end, state)


def _experience_values(state: CVAgentState) -> list[dict]:
    if not state.cv_draft:
        return []

    values = []
    for exp in state.cv_draft.experiences:
        bullets = exp.bullets or ([exp.raw_description] if exp.raw_description else [])
        company = compact_text(exp.company).upper()
        position = compact_text(exp.position)
        date = _date_range(exp.start_date, exp.end_date, state)
        values.append({
            "company": company,
            "company_with_date": f"{company}\t{date}",
            "position": position,
            "date": date,
            "bullets": compact_lines(bullets),
            "combined": f"{company}\t{date}",
        })
    return values


def _education_values(state: CVAgentState) -> list[dict]:
    if not state.cv_draft:
        return []

    values = []
    for edu in state.cv_draft.educations:
        degree_parts = [edu.degree, edu.major]
        degree = " - ".join(compact_text(part) for part in degree_parts if compact_text(part))
        school_name = compact_text(edu.school).upper()
        details = school_name
        if edu.gpa:
            details = f"{details}  GPA: {edu.gpa}"
        date = _date_range(edu.start_date, edu.end_date, state)
        values.append({
            "degree": degree,
            "school": details,
            "school_with_date": f"{school_name}\t{date}",
            "date": date,
            "combined": f"{school_name}\t{date}",
        })
    return values


def _skill_values(state: CVAgentState) -> list[str]:
    if not state.cv_draft or not state.cv_draft.skills_categorized:
        return state.user_profile.skills_raw if state.user_profile else []

    values = []
    for category, skills in state.cv_draft.skills_categorized.items():
        values.append(f"{category}: {', '.join(skills)}")
    return values


def _contact_parts(state: CVAgentState) -> dict:
    if not state.user_profile:
        return {"phone": "", "email": "", "location": "", "link": "", "line": ""}

    p = state.user_profile
    link = p.linkedin or p.github or ""
    parts = [p.email, p.phone, p.location]
    return {
        "phone": p.phone or "",
        "email": p.email or "",
        "location": p.location or "",
        "link": link,
        "line": " | ".join(part for part in parts if part),
    }


FIELD_BUILDERS = {
    "full_name": build_full_name,
    "job_title": build_job_title,
    "contact_line": build_contact_line,
    "email": lambda s: build_simple_field(s, "email"),
    "phone": lambda s: build_simple_field(s, "phone"),
    "location": lambda s: build_simple_field(s, "location"),
    "linkedin": lambda s: build_simple_field(s, "linkedin"),
    "github": lambda s: build_simple_field(s, "github"),
    "summary": build_summary,
    "profile": build_summary,
    "about_me": build_summary,
    "personal_summary": build_summary,
    "language": build_languages,
    "languages": build_languages,
    "reference": build_references,
    "references": build_references,
    "photo": build_photo_filename,
    "image": build_photo_filename,
}

SECTION_BUILDERS = {
    "experience_section": build_experience_lines,
    "skills_section": build_skills_lines,
    "education_section": build_education_lines,
    "language_section": build_language_lines,
    "reference_section": build_reference_lines,
}

SECTION_HEADERS = {
    "experience_section": "Experience",
    "skills_section": "Skills",
    "education_section": "Education",
    "language_section": "Languages",
    "reference_section": "References",
}

SECTION_HEADERS_VI = {
    "experience_section": "Kinh nghiệm làm việc",
    "skills_section": "Kỹ năng",
    "education_section": "Học vấn",
    "language_section": "Ngôn ngữ",
    "reference_section": "Tham chiếu",
}


def _fill_experience_blocks(paragraphs, blocks: list[tuple], experiences: list[dict]) -> int:
    filled = 0
    for block_index, (company_i, position_i, date_i, bullet_indices) in enumerate(blocks):
        if block_index >= len(experiences):
            indices = [i for i in (company_i, position_i, date_i) if i is not None] + list(bullet_indices)
            clear_xml_at(paragraphs, *indices)
            continue

        exp = experiences[block_index]
        if company_i is not None:
            set_xml_header_at(paragraphs, company_i, exp["company_with_date"])
        if position_i is not None:
            set_xml_text_at(paragraphs, position_i, exp["position"])
        if date_i is not None:
            set_xml_text_at(paragraphs, date_i, "")

        for i, para_i in enumerate(bullet_indices):
            set_xml_text_at(paragraphs, para_i, exp["bullets"][i] if i < len(exp["bullets"]) else "")
        filled += 1
    return filled


def _fill_combined_experience_blocks(paragraphs, blocks: list[tuple], experiences: list[dict]) -> int:
    filled = 0
    for block_index, (header_i, bullet_indices) in enumerate(blocks):
        if block_index >= len(experiences):
            clear_xml_at(paragraphs, header_i, *bullet_indices)
            continue

        exp = experiences[block_index]
        set_xml_header_at(paragraphs, header_i, exp["combined"])
        for i, para_i in enumerate(bullet_indices):
            set_xml_text_at(paragraphs, para_i, exp["bullets"][i] if i < len(exp["bullets"]) else "")
        filled += 1
    return filled


def _fill_education_blocks(paragraphs, blocks: list[tuple], educations: list[dict]) -> int:
    filled = 0
    for block_index, block in enumerate(blocks):
        if block_index >= len(educations):
            clear_xml_at(paragraphs, *[i for i in block if i is not None])
            continue

        edu = educations[block_index]
        degree_i, school_i, date_i = block
        set_xml_text_at(paragraphs, degree_i, edu["degree"])
        if school_i is not None:
            set_xml_header_at(paragraphs, school_i, edu["school_with_date"])
        if date_i is not None:
            set_xml_text_at(paragraphs, date_i, "")
        filled += 1
    return filled


def _fill_education_combined_blocks(paragraphs, blocks: list[tuple], educations: list[dict]) -> int:
    filled = 0
    for block_index, (info_i, date_i) in enumerate(blocks):
        if block_index >= len(educations):
            clear_xml_at(paragraphs, *[i for i in (info_i, date_i) if i is not None])
            continue

        edu = educations[block_index]
        set_xml_header_at(paragraphs, info_i, edu["combined"])
        if date_i is not None:
            set_xml_text_at(paragraphs, date_i, "")
        filled += 1
    return filled


def _fill_list(paragraphs, indices: list[int], values: list[str], bold: bool = False) -> int:
    for i, para_i in enumerate(indices):
        text = values[i] if i < len(values) else ""
        if bold:
            set_xml_bold_text_at(paragraphs, para_i, text)
        else:
            set_xml_text_at(paragraphs, para_i, text)
    return min(len(indices), len(values))


def _append_missing_template_sections(doc: Document, state: CVAgentState, template_id: str):
    added = []
    language = state.user_input.get("output_language", "vi")
    headers = SECTION_HEADERS if language == "en" else SECTION_HEADERS_VI
    if template_id in {"4", "5"}:
        language_lines = build_language_lines(state)
        if language_lines:
            append_section(doc, "language_section", language_lines, language=language)
            added.append(headers["language_section"])
        reference_lines = build_reference_lines(state)
        if reference_lines:
            append_section(doc, "reference_section", reference_lines, language=language)
            added.append(headers["reference_section"])
    return added


def render_known_template(doc: Document, state: CVAgentState, template_id: str) -> dict:
    """
    Deterministic renderer for the five bundled templates in `template cv/`.
    These files are exemplar DOCX layouts without placeholders, so relying on LLM
    section detection is brittle. The mappings below target the sample paragraph
    positions in each template and preserve the original layout/styles.
    """
    paragraphs = xml_paragraphs(doc)
    contact = _contact_parts(state)
    experiences = _experience_values(state)
    educations = _education_values(state)
    skills = _skill_values(state)
    languages = state.user_profile.languages if state.user_profile else []
    filled = 0

    if template_id == "1":
        set_xml_text_at(paragraphs, 0, build_full_name(state).upper())
        set_xml_text_at(paragraphs, 1, build_job_title(state).upper())
        set_xml_text_at(paragraphs, 12, contact["phone"])
        set_xml_text_at(paragraphs, 14, contact["email"])
        set_xml_text_at(paragraphs, 20, contact["location"])
        set_xml_text_at(paragraphs, 22, contact["link"])
        set_xml_text_at(paragraphs, 16, build_summary(state))
        filled += _fill_experience_blocks(
            paragraphs,
            [(23, 24, 26, [31, 32, 33]), (36, 37, 38, [42, 43]), (46, 47, 48, [54, 55])],
            experiences,
        )
        filled += _fill_list(paragraphs, [30, 35, 40, 41, 45], skills, bold=True)
        set_xml_text_at(paragraphs, 53, ", ".join(languages))
        filled += _fill_education_combined_blocks(paragraphs, [(66, 69), (68, 73)], educations)
        reference = build_references(state)
        set_xml_text_at(paragraphs, 64, reference)
        set_xml_text_at(paragraphs, 65, "")

    elif template_id == "2":
        # Main body.
        filled += _fill_education_blocks(paragraphs, [(7, 8, 10), (14, 15, None)], educations)
        filled += _fill_experience_blocks(
            paragraphs,
            [(22, 21, 24, [26]), (30, 29, 32, [34]), (38, 37, 40, [42]), (47, 46, 49, [51])],
            experiences,
        )
        reference = build_references(state)
        set_xml_text_at(paragraphs, 57, reference)
        clear_xml_at(paragraphs, 58, 59, 60, 61, 62, 63, 64)

        # Sidebar text boxes may appear after the body paragraphs in the XML.
        replace_xml_paragraph_exact(paragraphs, "Lorna Alvarado", build_full_name(state), max_count=2)
        replace_xml_paragraph_exact(paragraphs, "Contact", "Contact", max_count=2)
        replace_xml_paragraph_exact(paragraphs, "hello@reallygreatsite.com", contact["email"])
        replace_xml_paragraph_exact(paragraphs, "123 Anywhere St., Any City, ST  12345", contact["location"])
        set_following_nonempty_after_heading(paragraphs, "About Me", [build_summary(state)], occurrence=0)
        set_following_nonempty_after_heading(paragraphs, "About Me", [build_summary(state)], occurrence=1)
        set_following_nonempty_after_heading(paragraphs, "Skills", skills[:6], occurrence=0, bold=True)
        set_following_nonempty_after_heading(paragraphs, "Skills", skills[:6], occurrence=1, bold=True)

    elif template_id == "3":
        set_xml_text_at(paragraphs, 0, build_full_name(state).upper())
        set_xml_text_at(paragraphs, 1, build_job_title(state).upper())
        set_xml_text_at(paragraphs, 10, contact["phone"])
        set_xml_text_at(paragraphs, 12, "  ".join(part for part in [contact["email"], contact["location"]] if part))
        set_xml_text_at(paragraphs, 13, contact["link"])
        clear_xml_at(paragraphs, 14)
        set_xml_text_at(paragraphs, 16, build_summary(state))
        filled += _fill_experience_blocks(
            paragraphs,
            [(23, 24, 25, [32, 33, 34]), (38, 39, 40, [44]), (53, 54, 56, [66])],
            experiences,
        )
        filled += _fill_list(paragraphs, [31, 36, 37, 42, 43], skills, bold=True)
        set_xml_text_at(paragraphs, 51, ", ".join(languages))
        set_xml_text_at(paragraphs, 65, build_references(state))
        clear_xml_at(paragraphs, 70, 71, 72)
        filled += _fill_education_blocks(paragraphs, [(73, 74, 75)], educations)

    elif template_id == "4":
        set_xml_text_at(paragraphs, 0, build_full_name(state).upper())
        set_xml_text_at(paragraphs, 3, build_job_title(state))
        set_xml_text_at(paragraphs, 7, contact["line"])
        set_xml_text_at(paragraphs, 10, build_summary(state))
        filled += _fill_experience_blocks(
            paragraphs,
            [(15, 14, 18, [20, 21, 22]), (26, 25, 27, [29, 30, 31]), (35, 34, 36, [38, 39, 40])],
            experiences,
        )
        filled += _fill_education_blocks(paragraphs, [(44, 45, 46), (49, 50, 52)], educations)
        filled += _fill_list(paragraphs, [57, 58, 61, 62], skills, bold=True)

    elif template_id == "5":
        set_xml_text_at(paragraphs, 0, build_full_name(state).upper())
        set_xml_text_at(paragraphs, 1, build_job_title(state))
        set_xml_text_at(paragraphs, 3, contact["line"])
        set_xml_text_at(paragraphs, 5, build_summary(state))
        filled += _fill_combined_experience_blocks(
            paragraphs,
            [(7, [8, 9, 10]), (12, [13, 14, 15]), (17, [18, 19])],
            experiences,
        )
        filled += _fill_education_blocks(paragraphs, [(23, 24, 31), (28, 29, 36)], educations)
        filled += _fill_list(paragraphs, [42, 43, 44, 45, 46, 47, 48, 49], skills, bold=True)

    added = _append_missing_template_sections(doc, state, template_id)
    return {"filled": filled, "deleted": 0, "added": added}


def detect_template_type(doc: Document) -> str:
    for para in iter_all_paragraphs(doc):
        if PLACEHOLDER_PATTERN.search(para.text):
            return "placeholder"
    return "exemplar"


def remove_paragraph(paragraph):
    p_element = paragraph._element
    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)


def clear_paragraph_text(paragraph):
    for run in paragraph.runs:
        run.text = ""


def set_paragraph_text(paragraph, text: str, bold: Optional[bool] = None):
    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run("")

    run.text = text
    if bold is not None:
        run.bold = bold


def xml_paragraphs(doc: Document):
    """Return all XML paragraphs, including paragraphs inside text boxes/shapes."""
    return list(doc.part.element.findall(".//" + qn("w:p")))


def xml_paragraph_text(paragraph_element) -> str:
    return "".join(t.text or "" for t in paragraph_element.findall(".//" + qn("w:t")))


def _paragraph_properties(paragraph_element):
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    return p_pr


def _first_run(paragraph_element):
    run = paragraph_element.find(qn("w:r"))
    if run is None:
        run = OxmlElement("w:r")
        paragraph_element.append(run)
    return run


def _run_properties(run):
    r_pr = run.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run.insert(0, r_pr)
    return r_pr


def _set_run_bold(run, bold: bool):
    r_pr = _run_properties(run)
    for node in list(r_pr.findall(qn("w:b"))):
        r_pr.remove(node)
    if bold:
        r_pr.append(OxmlElement("w:b"))


def _set_run_size(run, points: int):
    r_pr = _run_properties(run)
    for node in list(r_pr.findall(qn("w:sz"))):
        r_pr.remove(node)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(points * 2))
    r_pr.append(size)


def _set_right_tab_stop(paragraph_element):
    p_pr = _paragraph_properties(paragraph_element)
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)

    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)


def _set_paragraph_keep_together(paragraph_element):
    p_pr = _paragraph_properties(paragraph_element)
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_xml_header_with_date(paragraph_element, text: str):
    left, _, right = text.partition("\t")
    run = _first_run(paragraph_element)
    r_pr = copy.deepcopy(run.find(qn("w:rPr")))

    for child in list(run):
        if child.tag != qn("w:rPr"):
            run.remove(child)

    if r_pr is not None and run.find(qn("w:rPr")) is None:
        run.insert(0, r_pr)

    for extra_run in paragraph_element.findall(qn("w:r")):
        if extra_run is run:
            continue
        for child in list(extra_run):
            if child.tag != qn("w:rPr"):
                extra_run.remove(child)

    _set_run_bold(run, True)
    _set_run_size(run, 12)
    _set_right_tab_stop(paragraph_element)
    _set_paragraph_keep_together(paragraph_element)

    left_node = OxmlElement("w:t")
    left_node.set(qn("xml:space"), "preserve")
    left_node.text = compact_text(left)
    run.append(left_node)

    if right.strip():
        run.append(OxmlElement("w:tab"))
        right_node = OxmlElement("w:t")
        right_node.set(qn("xml:space"), "preserve")
        right_node.text = compact_text(right)
        run.append(right_node)


def set_xml_header_at(paragraphs, index: int, text: str) -> bool:
    if 0 <= index < len(paragraphs):
        set_xml_header_with_date(paragraphs[index], text)
        return True
    return False


def set_xml_paragraph_text(paragraph_element, text: str):
    text_nodes = paragraph_element.findall(".//" + qn("w:t"))
    if not text_nodes:
        return

    text_nodes[0].text = compact_text(text)
    text_nodes[0].set(qn("xml:space"), "preserve")
    for node in text_nodes[1:]:
        node.text = ""


def set_xml_text_at(paragraphs, index: int, text: str) -> bool:
    if 0 <= index < len(paragraphs):
        set_xml_paragraph_text(paragraphs[index], text)
        return True
    return False


def set_xml_bold_text_at(paragraphs, index: int, text: str) -> bool:
    if not (0 <= index < len(paragraphs)):
        return False
    p = paragraphs[index]
    set_xml_paragraph_text(p, text)
    for run in p.findall(qn("w:r")):
        _set_run_bold(run, True)
    return True


def clear_xml_at(paragraphs, *indices: int):
    for index in indices:
        set_xml_text_at(paragraphs, index, "")


def set_following_nonempty_after_heading(paragraphs, heading: str, values: list[str], occurrence: int = 0, bold: bool = False) -> int:
    seen = 0
    for i, p in enumerate(paragraphs):
        if xml_paragraph_text(p).strip().lower() != heading.lower():
            continue
        if seen != occurrence:
            seen += 1
            continue

        changed = 0
        cursor = i + 1
        for value in values:
            while cursor < len(paragraphs) and not xml_paragraph_text(paragraphs[cursor]).strip():
                cursor += 1
            if cursor >= len(paragraphs):
                break
            set_xml_paragraph_text(paragraphs[cursor], value)
            if bold:
                for run in paragraphs[cursor].findall(qn("w:r")):
                    _set_run_bold(run, True)
            cursor += 1
            changed += 1
        return changed
    return 0


def replace_xml_paragraph_exact(paragraphs, old_text: str, new_text: str, max_count: Optional[int] = None) -> int:
    changed = 0
    for p in paragraphs:
        if xml_paragraph_text(p).strip() != old_text:
            continue
        set_xml_paragraph_text(p, new_text)
        changed += 1
        if max_count is not None and changed >= max_count:
            break
    return changed


def cleanup_consecutive_empty_paragraphs(doc: Document) -> int:
    """Collapse repeated empty paragraphs left by unused sample blocks."""
    removed = 0
    previous_empty = False
    for paragraph in list(doc.paragraphs):
        is_empty = not paragraph.text.strip()
        if is_empty and previous_empty:
            remove_paragraph(paragraph)
            removed += 1
            continue
        previous_empty = is_empty
    return removed


def insert_paragraph_after(paragraph, text: str, bold: bool = False):
    new_p_xml = copy.deepcopy(paragraph._element)
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for t_el in new_p_xml.findall(f'.//{ns}t'):
        t_el.text = ""

    paragraph._element.addnext(new_p_xml)
    new_para = paragraph.__class__(new_p_xml, paragraph._parent)
    set_paragraph_text(new_para, text, bold=bold)
    return new_para


def clone_paragraph_after(anchor, template_paragraph):
    new_p_xml = copy.deepcopy(template_paragraph._element)
    anchor_element = getattr(anchor, "_element", None)
    if anchor_element is None:
        anchor_element = anchor._tbl
    anchor_element.addnext(new_p_xml)
    return template_paragraph.__class__(new_p_xml, template_paragraph._parent)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def clear_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is not None:
        tc_pr.remove(margins)
    margins = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), "0")
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tc_pr.append(margins)


def set_cell_width(cell, width_twips: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def insert_company_time_table_after(paragraph, company: str, time_value: str):
    try:
        table = paragraph._parent.add_table(rows=1, cols=2, width=Inches(6.5))
    except TypeError:
        table = paragraph._parent.add_table(rows=1, cols=2)
    paragraph._element.addnext(table._tbl)

    remove_table_borders(table)
    table.autofit = False
    left_cell, right_cell = table.rows[0].cells
    set_cell_width(left_cell, 5200)
    set_cell_width(right_cell, 3800)
    clear_cell_margins(left_cell)
    clear_cell_margins(right_cell)

    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run(compact_text(company))
    left_run.bold = True

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.add_run(compact_text(time_value))

    remove_paragraph(paragraph)
    return table


def set_company_time_paragraph(paragraph, company: str, time_value: str):
    return insert_company_time_table_after(paragraph, company, time_value)


def set_bullet_paragraph(paragraph, text: str):
    value = compact_text(text).lstrip("•-–* ").strip()
    set_paragraph_text(paragraph, f"• {value}" if value else "")


def find_next_content_paragraph(paragraphs: list, start_index: int):
    for index in range(start_index, len(paragraphs)):
        if paragraphs[index].text.strip():
            return index, paragraphs[index]
    return None, None


def render_experience_template_blocks(doc: Document, state: CVAgentState) -> int:
    if not state.cv_draft:
        return 0

    rendered = 0
    paragraphs = list(doc.paragraphs)
    for index in range(len(paragraphs) - 1, -1, -1):
        header_para = paragraphs[index]
        header_text = header_para.text
        if "{{COMPANY}}" not in header_text or "{{TIME}}" not in header_text:
            continue

        bullet_index, bullet_para = find_next_content_paragraph(paragraphs, index + 1)
        if bullet_para is None or not (
            "{{EXPERIENCE}}" in bullet_para.text or "{{WORK_EXPERIENCE}}" in bullet_para.text
        ):
            continue

        experiences = state.cv_draft.experiences
        if not experiences:
            remove_paragraph(bullet_para)
            remove_paragraph(header_para)
            continue

        cursor = bullet_para
        for exp_index, exp in enumerate(experiences):
            current_header = header_para if exp_index == 0 else clone_paragraph_after(cursor, header_para)
            if exp_index != 0:
                cursor = current_header
            time_value = build_time_placeholder(exp.start_date, exp.end_date, state)
            cursor = set_company_time_paragraph(current_header, exp.company, time_value)

            bullets = compact_lines(exp.bullets or ([exp.raw_description] if exp.raw_description else []))
            if not bullets:
                bullets = [exp.position]
            for bullet_index, bullet in enumerate(bullets):
                current_bullet = bullet_para if exp_index == 0 and bullet_index == 0 else clone_paragraph_after(cursor, bullet_para)
                set_bullet_paragraph(current_bullet, bullet)
                cursor = current_bullet
        rendered += 1

    return rendered


def render_education_template_blocks(doc: Document, state: CVAgentState) -> int:
    if not state.cv_draft:
        return 0

    rendered = 0
    paragraphs = list(doc.paragraphs)
    for index in range(len(paragraphs) - 1, -1, -1):
        header_para = paragraphs[index]
        header_text = header_para.text
        if "{{SCHOOL}}" not in header_text or "{{TIME}}" not in header_text:
            continue

        detail_index, detail_para = find_next_content_paragraph(paragraphs, index + 1)
        if detail_para is None or "{{EDUCATION}}" not in detail_para.text:
            continue

        educations = state.cv_draft.educations
        if not educations:
            remove_paragraph(detail_para)
            remove_paragraph(header_para)
            continue

        cursor = detail_para
        for edu_index, edu in enumerate(educations):
            current_header = header_para if edu_index == 0 else clone_paragraph_after(cursor, header_para)
            if edu_index != 0:
                cursor = current_header
            time_value = build_time_placeholder(edu.start_date, edu.end_date, state)
            cursor = set_company_time_paragraph(current_header, edu.school, time_value)

            degree_parts = [compact_text(edu.degree), compact_text(edu.major)]
            detail = " - ".join(part for part in degree_parts if part)
            if edu.gpa:
                detail = f"{detail} | GPA: {edu.gpa}" if detail else f"GPA: {edu.gpa}"
            current_detail = detail_para if edu_index == 0 else clone_paragraph_after(cursor, detail_para)
            set_paragraph_text(current_detail, detail)
            cursor = current_detail
            for achievement in compact_lines(edu.achievements):
                current_achievement = clone_paragraph_after(cursor, detail_para)
                set_bullet_paragraph(current_achievement, achievement)
                cursor = current_achievement
        rendered += 1

    return rendered


# =====================================================================
# STRATEGY A: PLACEHOLDER TEMPLATE
# =====================================================================

class PlaceholderMapping(BaseModel):
    placeholder: str
    cv_field: str = Field(
        description=(
            "Field trong CV draft. Phải là MỘT trong: "
            "full_name, job_title, contact_line, email, phone, location, linkedin, github, "
            "summary, profile, about_me, personal_summary, language, reference, photo, "
            "experience_section, skills_section, education_section, language_section, "
            "reference_section, UNKNOWN"
        )
    )


class PlaceholderAnalysis(BaseModel):
    mappings: list[PlaceholderMapping]


PLACEHOLDER_PROMPT = """Map mỗi placeholder với field tương ứng.

Field hợp lệ:
- full_name, job_title, contact_line, email, phone, location, linkedin, github
- summary, profile, about_me, personal_summary, language, reference, photo
- experience_section, skills_section, education_section, language_section, reference_section
- UNKNOWN nếu không khớp

Quy tắc:
- {{NAME}}, {{FULLNAME}}, {{HO_TEN}} → full_name
- {{ABOUT}}, {{PROFILE}}, {{OBJECTIVE}}, {{PERSONAL_SUMMARY}} → summary
- {{LANGUAGE}}, {{LANGUAGES}} → language hoặc language_section
- {{REFERENCE}}, {{REFERENCES}} → reference hoặc reference_section
- {{PHOTO}}, {{IMAGE}}, {{AVATAR}} → photo
- {{WORK}}, {{JOB}}, {{CAREER}}, {{EXPERIENCE}} → experience_section"""


def render_placeholder_template(doc: Document, state: CVAgentState, llm) -> dict:
    block_fills = 0
    block_fills += render_experience_template_blocks(doc, state)
    block_fills += render_education_template_blocks(doc, state)

    placeholders = set()
    all_paragraphs = list(iter_all_paragraphs(doc))
    for para in all_paragraphs:
        for m in PLACEHOLDER_PATTERN.finditer(para.text):
            placeholders.add(m.group(1))

    if not placeholders:
        return {"filled": 0, "deleted": 0, "added": []}

    mapping = {ph: PLACEHOLDER_FIELD_MAP[ph] for ph in placeholders if ph in PLACEHOLDER_FIELD_MAP}
    unknown_placeholders = sorted(ph for ph in placeholders if ph not in mapping)
    if unknown_placeholders:
        user_msg = f"Placeholders: {unknown_placeholders}"
        analysis = call_with_structured_output(llm, PlaceholderAnalysis, PLACEHOLDER_PROMPT, user_msg)
        mapping.update({m.placeholder: m.cv_field for m in analysis.mappings})

    filled, deleted = block_fills, 0
    for ph in placeholders:
        cv_field = mapping.get(ph, "UNKNOWN")
        pattern = "{{" + ph + "}}"

        if cv_field in {"photo", "image"}:
            inserted = insert_photo_at_placeholder(doc, state, pattern)
            if not inserted:
                for para in all_paragraphs:
                    if pattern in para.text:
                        new_text = para.text.replace(pattern, "")
                        set_paragraph_text(para, new_text)
            if inserted:
                filled += 1
            else:
                deleted += 1

        elif cv_field in FIELD_BUILDERS:
            value = FIELD_BUILDERS[cv_field](state)
            for para in all_paragraphs:
                if pattern in para.text:
                    new_text = para.text.replace(pattern, value)
                    set_paragraph_text(para, new_text)
            if value:
                filled += 1
            else:
                deleted += 1

        elif cv_field in SECTION_BUILDERS:
            lines = SECTION_BUILDERS[cv_field](state)
            text_block = "\n".join(t for _, t in lines if t)
            for para in all_paragraphs:
                if pattern in para.text:
                    new_text = para.text.replace(pattern, text_block)
                    set_paragraph_text(para, new_text)
            if lines:
                filled += 1
            else:
                deleted += 1
        else:
            for para in all_paragraphs:
                if pattern in para.text:
                    new_text = para.text.replace(pattern, "")
                    set_paragraph_text(para, new_text)
            deleted += 1

    return {"filled": filled, "deleted": deleted, "added": []}


# =====================================================================
# STRATEGY B: EXEMPLAR TEMPLATE
# =====================================================================

class SectionDetection(BaseModel):
    heading_paragraph_index: int = Field(description="Index paragraph chứa heading")
    heading_text: str
    cv_field: str = Field(
        description=(
            "Field tương ứng: summary, experience_section, skills_section, "
            "education_section, language_section, reference_section, UNKNOWN"
        )
    )
    content_start_index: int
    content_end_index: int


class ExemplarAnalysis(BaseModel):
    sections: list[SectionDetection]
    name_paragraph_index: Optional[int] = None
    job_title_paragraph_index: Optional[int] = None
    contact_paragraph_index: Optional[int] = None


EXEMPLAR_PROMPT = """Bạn phân tích cấu trúc CV template để xác định các section.

Mỗi section gồm:
- Heading (bold, vd: "Summary", "Education", "Experience", "Skills")
- Content paragraphs sau heading, cho đến heading section tiếp theo

Map heading → cv_field:
- "Summary"/"About"/"About me"/"Profile"/"Personal Summary" → summary
- "Education" → education_section
- "Experience"/"Work Experience" → experience_section
- "Skills"/"Technical Skills" → skills_section
- "Language"/"Languages" → language_section
- "Reference"/"References" → reference_section
- Bỏ qua "Projects", "Certifications", "Leadership", "Activities" nếu template có các section này.
- Khác → UNKNOWN

QUAN TRỌNG:
- content_start_index = paragraph NGAY SAU heading
- content_end_index = paragraph CUỐI trước heading kế tiếp (inclusive)
- Index phải CHÍNH XÁC theo listing được cung cấp"""


def fill_section_content(doc: Document, section: SectionDetection, lines: list[tuple]) -> int:
    all_paras = doc.paragraphs
    start = section.content_start_index
    end = min(section.content_end_index, len(all_paras) - 1)

    if start > end or start >= len(all_paras):
        return 0

    content_paras = list(all_paras[start:end + 1])

    if not lines:
        for p in content_paras:
            remove_paragraph(p)
        return 0

    filtered_lines = [(t, compact_tabbed_text(txt)) for t, txt in lines if compact_text(txt)]

    para_idx = 0
    last_para = None

    for line_type, line_text in filtered_lines:
        is_bold = line_type in ("header", "subheader", "header_with_date")

        if para_idx < len(content_paras):
            p = content_paras[para_idx]
            if line_type == "header_with_date":
                left, _, right = line_text.partition("\t")
                clear_paragraph_text(p)
                run = p.runs[0] if p.runs else p.add_run("")
                run.text = compact_text(left)
                if right.strip():
                    run.add_tab()
                    run.add_text(compact_text(right))
                run.bold = True
                run.font.size = Pt(12)
                _set_right_tab_stop(p._element)
                _set_paragraph_keep_together(p._element)
            else:
                set_paragraph_text(p, line_text, bold=is_bold)
            last_para = p
            para_idx += 1
        else:
            if last_para is not None:
                last_para = insert_paragraph_after(last_para, line_text, bold=is_bold)

    while para_idx < len(content_paras):
        remove_paragraph(content_paras[para_idx])
        para_idx += 1

    return len(filtered_lines)


def append_section(doc: Document, section_field: str, lines: list[tuple], language: str = "vi"):
    if not lines:
        return

    headers = SECTION_HEADERS if language == "en" else SECTION_HEADERS_VI
    header_text = headers.get(section_field, section_field.replace("_", " ").title())

    doc.add_paragraph()

    h = doc.add_paragraph()
    h_run = h.add_run(header_text)
    h_run.bold = True
    h_run.font.size = Pt(13)

    for line_type, text in lines:
        if not text:
            continue
        p = doc.add_paragraph()
        run = p.add_run("")
        if line_type == "header_with_date":
            left, _, right = compact_tabbed_text(text).partition("\t")
            run.add_text(left)
            if right.strip():
                run.add_tab()
                run.add_text(compact_text(right))
            run.bold = True
            run.font.size = Pt(12)
            _set_right_tab_stop(p._element)
            _set_paragraph_keep_together(p._element)
        else:
            run.add_text(compact_text(text))
        if line_type in ("header", "subheader"):
            run.bold = True


def render_exemplar_template(doc: Document, state: CVAgentState, llm) -> dict:
    para_listing = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()[:120] if p.text.strip() else "(empty)"
        is_bold = bool(p.runs and p.runs[0].bold)
        bold_marker = "[BOLD] " if is_bold else ""
        para_listing.append(f"{i}: {bold_marker}{text}")

    listing_text = "\n".join(para_listing)
    user_msg = f"Phân tích template sau:\n\n{listing_text}\n\nXác định các section."

    analysis = call_with_structured_output(llm, ExemplarAnalysis, EXEMPLAR_PROMPT, user_msg)

    sections_sorted = sorted(
        analysis.sections,
        key=lambda s: s.heading_paragraph_index,
        reverse=True
    )

    sections_in_template = set()
    filled_count = 0
    deleted_count = 0

    for section in sections_sorted:
        if section.cv_field == "UNKNOWN":
            continue

        sections_in_template.add(section.cv_field)

        if section.cv_field in FIELD_BUILDERS:
            value = FIELD_BUILDERS[section.cv_field](state)
            if value:
                lines = [("text", value)]
                fill_section_content(doc, section, lines)
                filled_count += 1
            else:
                _delete_section_with_heading(doc, section)
                deleted_count += 1

        elif section.cv_field in SECTION_BUILDERS:
            lines = SECTION_BUILDERS[section.cv_field](state)
            if lines:
                fill_section_content(doc, section, lines)
                filled_count += 1
            else:
                _delete_section_with_heading(doc, section)
                deleted_count += 1

    header_fills = []
    if analysis.contact_paragraph_index is not None:
        header_fills.append((analysis.contact_paragraph_index, build_contact_line(state)))
    if analysis.job_title_paragraph_index is not None:
        header_fills.append((analysis.job_title_paragraph_index, build_job_title(state)))
    if analysis.name_paragraph_index is not None:
        header_fills.append((analysis.name_paragraph_index, build_full_name(state)))

    header_fills.sort(key=lambda x: x[0], reverse=True)

    for idx, value in header_fills:
        if not value:
            continue
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            template_run = p.runs[0] if p.runs else None
            was_bold = template_run.bold if template_run else None
            font_size = template_run.font.size if template_run else None

            clear_paragraph_text(p)
            run = p.runs[0] if p.runs else p.add_run("")
            run.text = value
            if was_bold is not None:
                run.bold = was_bold
            if font_size:
                run.font.size = font_size
            filled_count += 1

    added = []
    cv_has = set()
    if state.cv_draft:
        if state.cv_draft.experiences:
            cv_has.add("experience_section")
        if state.cv_draft.skills_categorized:
            cv_has.add("skills_section")
        if state.cv_draft.educations:
            cv_has.add("education_section")

    language = state.user_input.get("output_language", "vi")
    headers = SECTION_HEADERS if language == "en" else SECTION_HEADERS_VI
    missing = cv_has - sections_in_template
    for field in missing:
        if field in SECTION_BUILDERS:
            lines = SECTION_BUILDERS[field](state)
            if lines:
                append_section(doc, field, lines, language=language)
                added.append(headers.get(field, field))

    return {"filled": filled_count, "deleted": deleted_count, "added": added}


def _delete_section_with_heading(doc: Document, section: SectionDetection):
    all_paras = doc.paragraphs
    start = section.heading_paragraph_index
    end = min(section.content_end_index, len(all_paras) - 1)

    paras_to_delete = list(all_paras[start:end + 1])
    for p in reversed(paras_to_delete):
        remove_paragraph(p)


# =====================================================================
# MAIN NODE
# =====================================================================

def template_renderer_node(state: CVAgentState) -> dict:
    if not state.user_profile or not state.user_profile.template_path:
        return {"messages": ["[Template Renderer] Bỏ qua - không có template_path"]}

    template_path = state.user_profile.template_path
    p = Path(template_path)
    if not p.is_absolute():
        project_root = Path(settings.templates_dir).parent
        abs_template_path = str(project_root / template_path)
    else:
        abs_template_path = str(p)

    if not os.path.exists(abs_template_path):
        return {"messages": [
            f"[Template Renderer] LỖI: không tìm thấy template tại {abs_template_path}."
        ]}

    if not state.cv_draft:
        return {"messages": ["[Template Renderer] Không có CV draft để render"]}

    try:
        if Path(abs_template_path).suffix.lower() == ".html":
            from cvcraft.generate_cv.agents.html_renderer import fields_from_state, render_html_template

            outputs_dir = Path(settings.outputs_dir)
            outputs_dir.mkdir(parents=True, exist_ok=True)
            template_id = str(state.user_input.get("template_id") or Path(abs_template_path).stem)
            output_path = str(outputs_dir / f"cv_output_template_{template_id}.html")
            fields = fields_from_state(state)
            result = render_html_template(abs_template_path, output_path, fields)
            file_size = os.path.getsize(output_path)
            return {
                "output_path": output_path,
                "messages": [
                    f"[Template Renderer] (html-template) ĐÃ LƯU FILE: {output_path} "
                    f"({file_size} bytes). Fill {result['filled']}"
                ],
            }

        doc = Document(abs_template_path)
        template_id = str(state.user_input.get("template_id") or Path(abs_template_path).stem)
        template_kind = detect_template_type(doc)

        if template_kind == "placeholder":
            template_type = "placeholder"
            llm = LLMFactory.get_llm(tier="cheap")
            result = render_placeholder_template(doc, state, llm)
        elif template_id in KNOWN_TEMPLATE_IDS:
            template_type = f"known-{template_id}"
            result = render_known_template(doc, state, template_id)
            result["deleted"] += cleanup_consecutive_empty_paragraphs(doc)
        else:
            template_type = template_kind
            llm = LLMFactory.get_llm(
                tier="strong" if template_type == "exemplar" else "cheap"
            )

            result = render_exemplar_template(doc, state, llm)

        outputs_dir = Path(settings.outputs_dir)
        outputs_dir.mkdir(parents=True, exist_ok=True)
        output_path = str(outputs_dir / "cv_output.docx")

        doc.save(output_path)

        if not os.path.exists(output_path):
            return {"messages": [f"[Template Renderer] LỖI: file {output_path} không tồn tại sau khi save"]}

        file_size = os.path.getsize(output_path)
        added_msg = f", thêm: {result['added']}" if result["added"] else ""

        return {
            "output_path": output_path,
            "messages": [
                f"[Template Renderer] ({template_type}) ĐÃ LƯU FILE: {output_path} "
                f"({file_size} bytes). Fill {result['filled']}, xóa {result['deleted']}{added_msg}"
            ],
        }

    except Exception as e:
        import traceback
        return {"messages": [
            f"[Template Renderer] LỖI: {type(e).__name__}: {str(e)}\n"
            f"{traceback.format_exc()}"
        ]}
