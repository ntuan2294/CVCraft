"""
CV Parser — Extract raw text from an uploaded CV file.

Supported: PDF, DOCX, PNG, JPG/JPEG
For scanned PDFs or images: OpenAI vision API is used.
"""
import base64
import io
import os
from pathlib import Path


def parse_cv(data: bytes, mime: str, filename: str) -> str:
    """Return plain text extracted from a CV file."""
    ext = Path(filename).suffix.lower()

    if mime.startswith("image/") or ext in {".png", ".jpg", ".jpeg"}:
        return _parse_image(data, mime or "image/png")

    if ext == ".pdf" or mime == "application/pdf":
        text = _parse_pdf(data)
        if text.strip():
            return text
        # Scanned PDF — fall back to vision
        return _parse_image(data, "image/png")

    if ext == ".docx" or "wordprocessingml" in mime:
        return _parse_docx(data)

    raise ValueError(f"Định dạng không hỗ trợ: {ext or mime}")


def _parse_pdf(data: bytes) -> str:
    import fitz  # pymupdf — handles bold/special fonts better than pypdf
    doc = fitz.open(stream=data, filetype="pdf")
    parts = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(parts)


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _parse_image(data: bytes, mime: str) -> str:
    from openai import OpenAI
    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY chưa được cấu hình")

    b64 = base64.b64encode(data).decode()
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {
                        "type": "text",
                        "text": (
                            "Extract all text from this CV/resume image exactly as it appears. "
                            "Preserve structure: headings, bullet points, dates, contact info. "
                            "Return only the raw text content."
                        ),
                    },
                ],
            }
        ],
    )
    return resp.choices[0].message.content or ""
