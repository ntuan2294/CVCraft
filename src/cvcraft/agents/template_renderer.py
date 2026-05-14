"""
Template Renderer Agent - Layer 4.

Hỗ trợ 2 kiểu template:
A) Placeholder template: {{FIELD_NAME}} - dùng cho user kỹ thuật
B) Exemplar template: section header + nội dung mẫu

Auto-detect: nếu thấy {{...}} → kiểu A, không thì kiểu B.
"""
import os
import re
import copy
from pathlib import Path
from typing import Optional
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt
from cvcraft.core.state import CVAgentState
from cvcraft.infrastructure.llm.factory import LLMFactory, call_with_structured_output
from cvcraft.config.settings import settings


PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Z_][A-Z0-9_]*)\}\}')


def build_full_name(state: CVAgentState) -> str:
    return state.user_profile.full_name if state.user_profile else ""


def build_job_title(state: CVAgentState) -> str:
    if state.job_requirement and state.job_requirement.job_title:
        return state.job_requirement.job_title
    if state.user_profile and state.user_profile.work_experiences:
        return state.user_profile.work_experiences[0].position
    return ""


def build_contact_line(state: CVAgentState) -> str:
    if not state.user_profile:
        return ""
    p = state.user_profile
    parts = [p.location, p.email, p.phone]
    parts = [x for x in parts if x]
    return " • ".join(parts)


def build_simple_field(state: CVAgentState, field: str) -> str:
    if not state.user_profile:
        return ""
    return getattr(state.user_profile, field, "") or ""


def build_summary(state: CVAgentState) -> str:
    return state.cv_draft.summary if state.cv_draft and state.cv_draft.summary else ""


def build_experience_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.experiences:
        return []

    lines = []
    for exp in state.cv_draft.experiences:
        end = exp.end_date or "Present"
        lines.append(("header", exp.company))
        lines.append(("subheader", f"{exp.position}\t{exp.start_date} – {end}"))
        for bullet in exp.bullets:
            lines.append(("bullet", bullet))
        lines.append(("blank", ""))
    return lines


def build_skills_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.skills_categorized:
        return []
    return [("bullet", f"{cat}: {', '.join(skills)}")
            for cat, skills in state.cv_draft.skills_categorized.items()]


def build_education_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.educations:
        return []

    lines = []
    for edu in state.cv_draft.educations:
        end = edu.end_date or "Present"
        lines.append(("header", edu.school))
        lines.append(("subheader", f"{edu.degree} in {edu.major}\t{edu.start_date} – {end}"))
        if edu.gpa:
            lines.append(("bullet", f"GPA: {edu.gpa}"))
        for ach in edu.achievements:
            lines.append(("bullet", ach))
        lines.append(("blank", ""))
    return lines


def build_projects_lines(state: CVAgentState) -> list[tuple]:
    if not state.cv_draft or not state.cv_draft.projects:
        return []
    lines = []
    for proj in state.cv_draft.projects:
        if isinstance(proj, dict):
            name = proj.get("name", "Unnamed")
            desc = proj.get("description", "")
            lines.append(("header", name))
            if desc:
                lines.append(("bullet", desc))
            lines.append(("blank", ""))
    return lines


def build_leadership_lines(state: CVAgentState) -> list[tuple]:
    if state.cv_draft and state.cv_draft.projects:
        return build_projects_lines(state)
    return []


FIELD_BUILDERS = {
    "full_name": build_full_name,
    "job_title": build_job_title,
    "contact_line": build_contact_line,
    "email": lambda s: build_simple_field(s, "email"),
    "phone": lambda s: build_simple_field(s, "phone"),
    "location": lambda s: build_simple_field(s, "location"),
    "linkedin": lambda s: build_simple_field(s, "linkedin"),
    "github": lambda s: build_simple_field(s, "github"),
    "summary": build_summary,
}

SECTION_BUILDERS = {
    "experience_section": build_experience_lines,
    "skills_section": build_skills_lines,
    "education_section": build_education_lines,
    "projects_section": build_projects_lines,
    "leadership_section": build_leadership_lines,
}

SECTION_HEADERS = {
    "experience_section": "Experience",
    "skills_section": "Skills",
    "education_section": "Education",
    "projects_section": "Projects",
    "leadership_section": "Leadership & Activities",
}


def detect_template_type(doc: Document) -> str:
    for para in doc.paragraphs:
        if PLACEHOLDER_PATTERN.search(para.text):
            return "placeholder"
    return "exemplar"


def remove_paragraph(paragraph):
    p_element = paragraph._element
    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)


def clear_paragraph_text(paragraph):
    for run in paragraph.runs:
        run.text = ""


def set_paragraph_text(paragraph, text: str, bold: Optional[bool] = None):
    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run("")

    run.text = text
    if bold is not None:
        run.bold = bold


def insert_paragraph_after(paragraph, text: str, bold: bool = False):
    new_p_xml = copy.deepcopy(paragraph._element)
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for t_el in new_p_xml.findall(f'.//{ns}t'):
        t_el.text = ""

    paragraph._element.addnext(new_p_xml)
    new_para = paragraph.__class__(new_p_xml, paragraph._parent)
    set_paragraph_text(new_para, text, bold=bold)
    return new_para


# =====================================================================
# STRATEGY A: PLACEHOLDER TEMPLATE
# =====================================================================

class PlaceholderMapping(BaseModel):
    placeholder: str
    cv_field: str = Field(
        description=(
            "Field trong CV draft. Phải là MỘT trong: "
            "full_name, job_title, contact_line, email, phone, location, linkedin, github, "
            "summary, experience_section, skills_section, education_section, "
            "projects_section, UNKNOWN"
        )
    )


class PlaceholderAnalysis(BaseModel):
    mappings: list[PlaceholderMapping]


PLACEHOLDER_PROMPT = """Map mỗi placeholder với field tương ứng.

Field hợp lệ:
- full_name, job_title, contact_line, email, phone, location, linkedin, github
- summary, experience_section, skills_section, education_section, projects_section
- UNKNOWN nếu không khớp

Quy tắc:
- {{NAME}}, {{FULLNAME}}, {{HO_TEN}} → full_name
- {{ABOUT}}, {{PROFILE}}, {{OBJECTIVE}} → summary
- {{WORK}}, {{JOB}}, {{CAREER}}, {{EXPERIENCE}} → experience_section"""


def render_placeholder_template(doc: Document, state: CVAgentState, llm) -> dict:
    placeholders = set()
    for para in doc.paragraphs:
        for m in PLACEHOLDER_PATTERN.finditer(para.text):
            placeholders.add(m.group(1))

    if not placeholders:
        return {"filled": 0, "deleted": 0, "added": []}

    user_msg = f"Placeholders: {sorted(placeholders)}"
    analysis = call_with_structured_output(llm, PlaceholderAnalysis, PLACEHOLDER_PROMPT, user_msg)
    mapping = {m.placeholder: m.cv_field for m in analysis.mappings}

    filled, deleted = 0, 0
    for ph in placeholders:
        cv_field = mapping.get(ph, "UNKNOWN")
        pattern = "{{" + ph + "}}"

        if cv_field in FIELD_BUILDERS:
            value = FIELD_BUILDERS[cv_field](state)
            for para in doc.paragraphs:
                if pattern in para.text:
                    new_text = para.text.replace(pattern, value)
                    set_paragraph_text(para, new_text)
            if value:
                filled += 1
            else:
                deleted += 1

        elif cv_field in SECTION_BUILDERS:
            lines = SECTION_BUILDERS[cv_field](state)
            text_block = "\n".join(t for _, t in lines if t)
            for para in doc.paragraphs:
                if pattern in para.text:
                    new_text = para.text.replace(pattern, text_block)
                    set_paragraph_text(para, new_text)
            if lines:
                filled += 1
            else:
                deleted += 1
        else:
            for para in doc.paragraphs:
                if pattern in para.text:
                    new_text = para.text.replace(pattern, "")
                    set_paragraph_text(para, new_text)
            deleted += 1

    return {"filled": filled, "deleted": deleted, "added": []}


# =====================================================================
# STRATEGY B: EXEMPLAR TEMPLATE
# =====================================================================

class SectionDetection(BaseModel):
    heading_paragraph_index: int = Field(
        description="Index paragraph chứa heading"
    )
    heading_text: str
    cv_field: str = Field(
        description=(
            "Field tương ứng: summary, experience_section, skills_section, "
            "education_section, projects_section, leadership_section, UNKNOWN"
        )
    )
    content_start_index: int
    content_end_index: int


class ExemplarAnalysis(BaseModel):
    sections: list[SectionDetection]
    name_paragraph_index: Optional[int] = None
    job_title_paragraph_index: Optional[int] = None
    contact_paragraph_index: Optional[int] = None


EXEMPLAR_PROMPT = """Bạn phân tích cấu trúc CV template để xác định các section.

Mỗi section gồm:
- Heading (bold, vd: "Summary", "Education", "Experience", "Skills")
- Content paragraphs sau heading, cho đến heading section tiếp theo

Map heading → cv_field:
- "Summary"/"About"/"Profile" → summary
- "Education" → education_section
- "Experience"/"Work Experience" → experience_section
- "Skills"/"Technical Skills" → skills_section
- "Projects" → projects_section
- "Leadership"/"Activities" → leadership_section
- Khác → UNKNOWN

QUAN TRỌNG:
- content_start_index = paragraph NGAY SAU heading
- content_end_index = paragraph CUỐI trước heading kế tiếp (inclusive)
- Index phải CHÍNH XÁC theo listing được cung cấp"""


def fill_section_content(doc: Document, section: SectionDetection, lines: list[tuple]) -> int:
    all_paras = doc.paragraphs
    start = section.content_start_index
    end = min(section.content_end_index, len(all_paras) - 1)

    if start > end or start >= len(all_paras):
        return 0

    content_paras = list(all_paras[start:end + 1])

    if not lines:
        for p in content_paras:
            remove_paragraph(p)
        return 0

    filtered_lines = [(t, txt) for t, txt in lines if txt or t == "blank"]

    para_idx = 0
    last_para = None

    for line_type, line_text in filtered_lines:
        if line_type == "blank":
            if para_idx < len(content_paras):
                p = content_paras[para_idx]
                clear_paragraph_text(p)
                last_para = p
                para_idx += 1
            continue

        is_bold = line_type in ("header", "subheader")

        if para_idx < len(content_paras):
            p = content_paras[para_idx]
            set_paragraph_text(p, line_text, bold=is_bold)
            last_para = p
            para_idx += 1
        else:
            if last_para is not None:
                last_para = insert_paragraph_after(last_para, line_text, bold=is_bold)

    while para_idx < len(content_paras):
        remove_paragraph(content_paras[para_idx])
        para_idx += 1

    return len([l for l in filtered_lines if l[1]])


def append_section(doc: Document, section_field: str, lines: list[tuple]):
    if not lines:
        return

    header_text = SECTION_HEADERS.get(section_field, section_field.replace("_", " ").title())

    doc.add_paragraph()

    h = doc.add_paragraph()
    h_run = h.add_run(header_text)
    h_run.bold = True
    h_run.font.size = Pt(13)

    for line_type, text in lines:
        if not text:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        run = p.add_run(text)
        if line_type in ("header", "subheader"):
            run.bold = True


def render_exemplar_template(doc: Document, state: CVAgentState, llm) -> dict:
    para_listing = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()[:120] if p.text.strip() else "(empty)"
        is_bold = bool(p.runs and p.runs[0].bold)
        bold_marker = "[BOLD] " if is_bold else ""
        para_listing.append(f"{i}: {bold_marker}{text}")

    listing_text = "\n".join(para_listing)
    user_msg = f"Phân tích template sau:\n\n{listing_text}\n\nXác định các section."

    analysis = call_with_structured_output(llm, ExemplarAnalysis, EXEMPLAR_PROMPT, user_msg)

    sections_sorted = sorted(
        analysis.sections,
        key=lambda s: s.heading_paragraph_index,
        reverse=True
    )

    sections_in_template = set()
    filled_count = 0
    deleted_count = 0

    for section in sections_sorted:
        if section.cv_field == "UNKNOWN":
            continue

        sections_in_template.add(section.cv_field)

        if section.cv_field in FIELD_BUILDERS:
            value = FIELD_BUILDERS[section.cv_field](state)
            if value:
                lines = [("text", value)]
                fill_section_content(doc, section, lines)
                filled_count += 1
            else:
                _delete_section_with_heading(doc, section)
                deleted_count += 1

        elif section.cv_field in SECTION_BUILDERS:
            lines = SECTION_BUILDERS[section.cv_field](state)
            if lines:
                fill_section_content(doc, section, lines)
                filled_count += 1
            else:
                _delete_section_with_heading(doc, section)
                deleted_count += 1

    header_fills = []
    if analysis.contact_paragraph_index is not None:
        header_fills.append((analysis.contact_paragraph_index, build_contact_line(state)))
    if analysis.job_title_paragraph_index is not None:
        header_fills.append((analysis.job_title_paragraph_index, build_job_title(state)))
    if analysis.name_paragraph_index is not None:
        header_fills.append((analysis.name_paragraph_index, build_full_name(state)))

    header_fills.sort(key=lambda x: x[0], reverse=True)

    for idx, value in header_fills:
        if not value:
            continue
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            template_run = p.runs[0] if p.runs else None
            was_bold = template_run.bold if template_run else None
            font_size = template_run.font.size if template_run else None

            clear_paragraph_text(p)
            run = p.runs[0] if p.runs else p.add_run("")
            run.text = value
            if was_bold is not None:
                run.bold = was_bold
            if font_size:
                run.font.size = font_size
            filled_count += 1

    added = []
    cv_has = set()
    if state.cv_draft:
        if state.cv_draft.experiences:
            cv_has.add("experience_section")
        if state.cv_draft.skills_categorized:
            cv_has.add("skills_section")
        if state.cv_draft.educations:
            cv_has.add("education_section")
        if state.cv_draft.projects:
            cv_has.add("projects_section")

    missing = cv_has - sections_in_template
    for field in missing:
        if field in SECTION_BUILDERS:
            lines = SECTION_BUILDERS[field](state)
            if lines:
                append_section(doc, field, lines)
                added.append(SECTION_HEADERS.get(field, field))

    return {"filled": filled_count, "deleted": deleted_count, "added": added}


def _delete_section_with_heading(doc: Document, section: SectionDetection):
    all_paras = doc.paragraphs
    start = section.heading_paragraph_index
    end = min(section.content_end_index, len(all_paras) - 1)

    paras_to_delete = list(all_paras[start:end + 1])
    for p in reversed(paras_to_delete):
        remove_paragraph(p)


# =====================================================================
# MAIN NODE
# =====================================================================

def template_renderer_node(state: CVAgentState) -> dict:
    if not state.user_profile or not state.user_profile.template_path:
        return {"messages": ["[Template Renderer] Bỏ qua - không có template_path"]}

    template_path = state.user_profile.template_path
    abs_template_path = os.path.abspath(template_path)

    if not os.path.exists(abs_template_path):
        return {"messages": [
            f"[Template Renderer] LỖI: không tìm thấy template tại {abs_template_path}."
        ]}

    if not state.cv_draft:
        return {"messages": ["[Template Renderer] Không có CV draft để render"]}

    try:
        doc = Document(abs_template_path)
        template_type = detect_template_type(doc)

        llm = LLMFactory.get_llm(
            tier="strong" if template_type == "exemplar" else "cheap"
        )

        if template_type == "placeholder":
            result = render_placeholder_template(doc, state, llm)
        else:
            result = render_exemplar_template(doc, state, llm)

        # Lưu vào outputs/ directory
        outputs_dir = Path(settings.outputs_dir)
        outputs_dir.mkdir(parents=True, exist_ok=True)
        output_path = str(outputs_dir / "cv_output.docx")

        doc.save(output_path)

        if not os.path.exists(output_path):
            return {"messages": [f"[Template Renderer] LỖI: file {output_path} không tồn tại sau khi save"]}

        file_size = os.path.getsize(output_path)
        added_msg = f", thêm: {result['added']}" if result["added"] else ""

        return {
            "output_path": output_path,
            "messages": [
                f"[Template Renderer] ({template_type}) ĐÃ LƯU FILE: {output_path} "
                f"({file_size} bytes). Fill {result['filled']}, xóa {result['deleted']}{added_msg}"
            ],
        }

    except Exception as e:
        import traceback
        return {"messages": [
            f"[Template Renderer] LỖI: {type(e).__name__}: {str(e)}\n"
            f"{traceback.format_exc()}"
        ]}
